from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import RFQExtract, SiteResearch, ProposalSections


def _set_default_styles(doc: Document) -> None:
    styles = doc.styles

    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"


def _add_title(doc: Document, rfq: RFQExtract) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("RAIN Consulting")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Draft Proposal")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    doc.add_heading(rfq.project_title or "Untitled Project", level=1)


def _add_project_details(doc: Document, rfq: RFQExtract, research: SiteResearch) -> None:
    doc.add_heading("Project Details", level=2)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Project type", rfq.project_type),
        ("Site address", research.address or rfq.site_address),
        ("Council", research.council),
        ("Traditional Owners", research.traditional_owners),
        ("Catchment Management Authority", research.cma),
        ("Water authority", research.water_authority),
        ("Latitude", research.latitude),
        ("Longitude", research.longitude),
    ]

    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value or "To be confirmed"


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_proposal_docx(
    output_path: str,
    rfq: RFQExtract,
    research: SiteResearch,
    sections: ProposalSections,
) -> None:
    doc = Document()
    _set_default_styles(doc)

    _add_title(doc, rfq)
    _add_project_details(doc, rfq, research)

    doc.add_heading("Project Understanding", level=2)
    for para in (sections.project_understanding or "").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading("Scope of Services", level=2)
    doc.add_paragraph(sections.scope_of_services or "To be confirmed.")

    doc.add_heading("Deliverables", level=2)
    _add_bullets(doc, sections.deliverables)

    doc.add_heading("Authority and Planning Considerations", level=2)
    if sections.authority_requirements:
        _add_bullets(doc, sections.authority_requirements)
    else:
        doc.add_paragraph("Authority requirements are to be confirmed.")

    doc.add_heading("Planning Controls", level=2)
    planning = research.planning
    _add_bullets(
        doc,
        [
            f"Zone: {planning.zone or 'To be confirmed'}",
            f"DPO: {planning.dpo or 'To be confirmed'}",
            f"SBO: {planning.sbo or 'To be confirmed'}",
            f"LSIO: {planning.lsio or 'To be confirmed'}",
            f"FO: {planning.fo or 'To be confirmed'}",
        ],
    )

    if planning.notes:
        doc.add_heading("Planning Notes", level=3)
        _add_bullets(doc, planning.notes)

    doc.add_heading("Assumptions", level=2)
    _add_bullets(doc, sections.assumptions)

    doc.add_heading("Exclusions", level=2)
    _add_bullets(doc, sections.exclusions)

    if rfq.extraction_notes or research.notes:
        doc.add_heading("Internal Review Notes", level=2)
        _add_bullets(doc, rfq.extraction_notes + research.notes)

    doc.save(output_path)
